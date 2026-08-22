"""
Tests for the data ingestion layer: sample data seeder, document store,
and local file connectors (CSV/DOCX/PDF/TXT/JSON).
"""
import io
import os
import sqlite3
import tempfile

import pytest

from agent.connectors.sample_data import seed_sample_data, sample_data_available
from agent.connectors.local import LocalFileConnector
from agent.connectors.storage import document_store
from agent.connectors.base import Document


# ==================== SAMPLE DATA ====================

class TestSampleData:
    def test_seeds_demo_tables(self):
        fd, path = tempfile.mkstemp(suffix=".db")
        os.close(fd)
        try:
            result = seed_sample_data(f"sqlite:///{path}")
            assert result["seeded"] is True
            assert set(result["tables"]) == {"customers", "products", "orders", "order_items"}

            conn = sqlite3.connect(path)
            counts = {
                t: conn.execute(f"SELECT COUNT(*) FROM {t}").fetchone()[0]
                for t in ("customers", "products", "orders", "order_items")
            }
            conn.close()
            assert counts["customers"] == 60
            assert counts["products"] == 8
            assert counts["orders"] == 1200
            assert counts["order_items"] == 1200
        finally:
            os.unlink(path)

    def test_idempotent(self):
        fd, path = tempfile.mkstemp(suffix=".db")
        os.close(fd)
        try:
            url = f"sqlite:///{path}"
            seed_sample_data(url)
            result = seed_sample_data(url)
            assert result["seeded"] is False  # already present
            assert sample_data_available(url) is True
        finally:
            os.unlink(path)

    def test_force_reseeds(self):
        fd, path = tempfile.mkstemp(suffix=".db")
        os.close(fd)
        try:
            url = f"sqlite:///{path}"
            seed_sample_data(url)
            result = seed_sample_data(url, force=True)
            assert result["seeded"] is True
        finally:
            os.unlink(path)

    def test_deterministic(self):
        """Same seed -> same data, so tests and demos are reproducible."""
        fd1, p1 = tempfile.mkstemp(suffix=".db")
        fd2, p2 = tempfile.mkstemp(suffix=".db")
        os.close(fd1)
        os.close(fd2)
        try:
            seed_sample_data(f"sqlite:///{p1}")
            seed_sample_data(f"sqlite:///{p2}")
            c1 = sqlite3.connect(p1)
            c2 = sqlite3.connect(p2)
            r1 = c1.execute("SELECT SUM(quantity * unit_price) FROM orders").fetchone()
            r2 = c2.execute("SELECT SUM(quantity * unit_price) FROM orders").fetchone()
            c1.close()
            c2.close()
            assert r1 == r2
        finally:
            os.unlink(p1)
            os.unlink(p2)


# ==================== LOCAL FILE CONNECTOR ====================

class TestLocalFileConnector:
    def test_csv_to_markdown_table(self):
        connector = LocalFileConnector()
        csv_bytes = b"name,revenue\nAcme,100\nGlobex,200\n"
        doc = connector.ingest_bytes("sales.csv", csv_bytes)
        assert doc.source == "local"
        assert doc.source_id == "sales.csv"
        assert "| name | revenue |" in doc.content
        assert "| Acme | 100 |" in doc.content
        assert doc.metadata["file_type"] == "csv"

    def test_tsv_supported(self):
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("data.tsv", b"a\tb\n1\t2\n")
        assert "| a | b |" in doc.content

    def test_txt_raw_text(self):
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("notes.txt", b"hello world")
        assert doc.content == "hello world"

    def test_json_pretty_printed(self):
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("config.json", b'{"a": 1}')
        assert '"a": 1' in doc.content

    def test_docx_parsed(self):
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")
        buf = io.BytesIO()
        d = DocxDocument()
        d.add_paragraph("Quarterly revenue grew 12%.")
        d.save(buf)
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("report.docx", buf.getvalue())
        assert "Quarterly revenue grew 12%." in doc.content

    def test_pdf_parsed(self):
        try:
            from pypdf import PdfWriter
        except ImportError:
            pytest.skip("pypdf not installed")
        buf = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        writer.write(buf)
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("blank.pdf", buf.getvalue())
        # pypdf parses the file; a blank page has no extractable text
        assert doc.content == "(no extractable text in PDF)"

    def test_unsupported_extension_rejected(self):
        connector = LocalFileConnector()
        with pytest.raises(ValueError, match="Unsupported file type"):
            connector.ingest_bytes("virus.exe", b"MZ")

    def test_oversized_file_rejected(self):
        connector = LocalFileConnector()
        with pytest.raises(ValueError, match="too large"):
            connector.ingest_bytes("big.csv", b"x" * (11 * 1024 * 1024))

    def test_empty_csv(self):
        connector = LocalFileConnector()
        doc = connector.ingest_bytes("empty.csv", b"")
        assert doc.content == "(empty file)"


# ==================== DOCUMENT STORE ====================

class TestDocumentStore:
    def test_upsert_dedupes_by_source_id(self):
        doc1 = Document(source="csv", source_id="a.csv", title="A", content="v1")
        doc2 = Document(source="csv", source_id="a.csv", title="A", content="v2")
        assert document_store.save_document(doc1) is True  # created
        assert document_store.save_document(doc2) is False  # updated
        docs = document_store.list_documents(source="csv")
        assert len(docs) == 1
        assert docs[0]["content"] == "v2"

    def test_list_and_count(self):
        document_store.save_document(Document(source="test-src", source_id="1", title="One", content="x"))
        document_store.save_document(Document(source="test-src", source_id="2", title="Two", content="y"))
        assert document_store.count_documents(source="test-src") == 2
        assert document_store.count_documents() >= 2
        docs = document_store.list_documents(source="test-src")
        assert {d["title"] for d in docs} == {"One", "Two"}

    def test_delete_document(self):
        document_store.save_document(Document(source="test-del", source_id="1", title="D", content="x"))
        docs = document_store.list_documents(source="test-del")
        assert document_store.delete_document(docs[0]["id"]) is True
        assert document_store.delete_document(999999) is False
        assert document_store.count_documents(source="test-del") == 0

    def test_delete_source(self):
        document_store.save_document(Document(source="test-src2", source_id="1", title="A", content="x"))
        document_store.save_document(Document(source="test-src2", source_id="2", title="B", content="y"))
        assert document_store.delete_source("test-src2") == 2

    def test_all_content_concatenates(self):
        document_store.save_document(Document(source="test-ctx", source_id="1", title="T", content="body text"))
        text = document_store.all_content()
        assert "body text" in text

    def test_connector_state_roundtrip(self):
        document_store.save_state("test-conn", config={"cursor": "abc"}, last_error=None)
        state = document_store.get_state("test-conn")
        assert state["config"] == {"cursor": "abc"}
        assert state["last_sync_at"] is not None
        document_store.save_state("test-conn", last_error="boom")
        assert document_store.get_state("test-conn")["last_error"] == "boom"
