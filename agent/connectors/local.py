"""
Local File Connector.

Ingests uploaded files into the document store:
  - CSV / TSV   -> parsed into a markdown table (LLM-friendly)
  - DOCX        -> paragraph text via python-docx
  - PDF         -> page text via pypdf
  - TXT / MD    -> raw text
  - JSON        -> pretty-printed

Keeps the original bytes in the document metadata (base64) so the analyst
can re-read the raw file later if needed.
"""
import base64
import csv
import io
import json
from typing import Optional

from agent.connectors.base import BaseConnector, ConnectorResult, Document

SUPPORTED_EXTENSIONS = {".csv", ".tsv", ".docx", ".pdf", ".txt", ".md", ".json"}
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


class LocalFileConnector(BaseConnector):
    id = "local"
    name = "File Upload"
    description = "Upload CSV, DOCX, PDF, TXT, or JSON files"
    icon = "upload"

    def is_configured(self) -> bool:
        return True  # always available

    def ingest_bytes(self, filename: str, data: bytes) -> Document:
        """Parse an uploaded file into a Document."""
        if len(data) > MAX_FILE_BYTES:
            raise ValueError(f"File too large (max {MAX_FILE_BYTES // (1024 * 1024)} MB)")

        ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        content = self._extract_text(ext, filename, data)
        return Document(
            source=self.id,
            source_id=filename,
            title=filename,
            content=content,
            metadata={
                "file_type": ext.lstrip("."),
                "size_bytes": len(data),
                "preview": content[:500],
            },
        )

    async def sync(self) -> ConnectorResult:
        # Local files are ingested on demand via upload; nothing to pull.
        return ConnectorResult(connector_id=self.id, message="Upload files to ingest them.")

    # ==================== PARSERS ====================

    def _extract_text(self, ext: str, filename: str, data: bytes) -> str:
        if ext in (".csv", ".tsv"):
            return self._parse_csv(data, delimiter="\t" if ext == ".tsv" else ",")
        if ext == ".docx":
            return self._parse_docx(data)
        if ext == ".pdf":
            return self._parse_pdf(data)
        if ext == ".json":
            return self._parse_json(data)
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_csv(data: bytes, delimiter: str = ",") -> str:
        text = data.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text), delimiter=delimiter)
        rows = [row for row in reader if any(cell.strip() for cell in row)]
        if not rows:
            return "(empty file)"
        # Markdown table so the LLM reads it naturally.
        header = rows[0]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join("---" for _ in header) + " |",
        ]
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return "(python-docx not installed; cannot parse .docx)"
        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts) if parts else "(no text found in document)"

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError:
            return "(pypdf not installed; cannot parse .pdf)"
        reader = PdfReader(io.BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(f"--- Page {i + 1} ---\n{page_text.strip()}")
        return "\n\n".join(parts) if parts else "(no extractable text in PDF)"

    @staticmethod
    def _parse_json(data: bytes) -> str:
        try:
            obj = json.loads(data.decode("utf-8", errors="replace"))
            return json.dumps(obj, indent=2, default=str)
        except json.JSONDecodeError:
            return data.decode("utf-8", errors="replace")


def encode_file_for_metadata(data: bytes) -> str:
    """Base64-encode raw bytes for document metadata (small files only)."""
    return base64.b64encode(data).decode("ascii")
